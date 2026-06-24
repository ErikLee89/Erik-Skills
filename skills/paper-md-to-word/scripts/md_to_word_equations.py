from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


ROOT = Path.cwd()
MML2OMML = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")
PANDOC = shutil.which("pandoc") or "pandoc"
EQUATION_MAP: dict[str, str] = {}


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5, bold=False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), east_asia)
    rpr.rFonts.set(qn("w:ascii"), ascii_font)
    rpr.rFonts.set(qn("w:hAnsi"), ascii_font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_borders(table, color="7F7F7F", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _next_word_id(elements, attr_name: str) -> int:
    values = []
    for element in elements:
        value = element.get(qn(attr_name))
        if value is not None and value.isdigit():
            values.append(int(value))
    return (max(values) + 1) if values else 1


def ensure_reference_numbering(doc) -> str:
    """Return a Word numbering id whose level text is '[%1]'."""
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        abstract_id = abstract.get(qn("w:abstractNumId"))
        for lvl_text in abstract.findall(".//" + qn("w:lvlText")):
            if lvl_text.get(qn("w:val")) == "[%1]":
                for num in numbering.findall(qn("w:num")):
                    abstract_num_id = num.find(qn("w:abstractNumId"))
                    if abstract_num_id is not None and abstract_num_id.get(qn("w:val")) == abstract_id:
                        return num.get(qn("w:numId"))

    abstract_id = str(_next_word_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId"))
    num_id = str(_next_word_id(numbering.findall(qn("w:num")), "w:numId"))

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "[%1]"), ("w:lvlJc", "left")):
        child = OxmlElement(tag)
        child.set(qn("w:val"), value)
        lvl.append(child)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "440")
    ind.set(qn("w:hanging"), "440")
    p_pr.append(ind)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:hint"), "eastAsia")
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), abstract_id)
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_reference_numbering(paragraph, num_id: str):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), num_id)
    num_pr.append(num)
    p_pr.insert(0, num_pr)



def clean_inline_math(text: str) -> str:
    def repl(match):
        value = match.group(1)
        value = value.replace(r"\boldsymbol", "")
        value = value.replace(r"\mathrm", "")
        value = value.replace(r"\left", "").replace(r"\right", "")
        value = value.replace(r"\,", " ")
        value = value.replace(r"\eta", "η").replace(r"\rho", "ρ").replace(r"\beta", "β")
        value = value.replace(r"^\circ", "°").replace(r"\circ", "°")
        value = value.replace("{", "").replace("}", "")
        value = value.replace(r"\pm", "±")
        value = value.replace(r"\times", "×")
        value = re.sub(r"_([A-Za-z0-9]+)", r"_\1", value)
        return value

    text = re.sub(r"\$(.+?)\$", repl, text)
    return text.replace("`", "")

def mathml_to_omml(mathml: str):
    transform = etree.XSLT(etree.parse(str(MML2OMML)))
    return transform(etree.fromstring(mathml.encode("utf-8"))).getroot()


def load_equation_map(path: Path | None) -> dict[str, str]:
    """Load optional LaTeX-to-MathML overrides from JSON.

    The JSON must be an object whose keys are LaTeX expressions and whose values
    are complete MathML <math>...</math> strings. This keeps project-specific
    formula handling outside the reusable script.
    """
    if path is None:
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Equation map JSON must be an object: {latex: mathml}.")
    return {normalize_equation(str(key)): str(value) for key, value in data.items()}


def split_equation_tag(latex: str) -> tuple[str, str | None]:
    match = re.search(r"\\tag\{([^{}]+)\}", latex)
    tag = match.group(1).strip() if match else None
    body = re.sub(r"\\tag\{[^{}]+\}", "", latex).strip()
    return body, tag


def pandoc_latex_to_mathml(latex: str, display: bool = False) -> str:
    delimiter = "$$" if display else "$"
    markdown = f"{delimiter}{latex}{delimiter}"
    command = [str(PANDOC), "-f", "markdown+tex_math_dollars", "-t", "html", "--mathml"]
    try:
        completed = subprocess.run(
            command,
            input=markdown,
            text=True,
            encoding="utf-8",
            capture_output=True,
            check=False,
        )
    except FileNotFoundError as exc:
        raise RuntimeError("Pandoc is required for generic LaTeX math conversion. Install pandoc or pass --pandoc.") from exc
    if completed.returncode != 0:
        raise RuntimeError(f"Pandoc math conversion failed for {latex!r}: {completed.stderr.strip()}")
    match = re.search(r"(<math\b.*?</math>)", completed.stdout, flags=re.S)
    if not match:
        message = completed.stderr.strip() or completed.stdout.strip()
        raise RuntimeError(f"Pandoc did not return MathML for {latex!r}: {message}")
    return match.group(1)


def latex_to_mathml(value: str, display: bool = False) -> str:
    value = normalize_equation(value.strip())
    value, _ = split_equation_tag(value)
    override = EQUATION_MAP.get(normalize_equation(value))
    if override:
        return override
    return pandoc_latex_to_mathml(value, display=display)


def inline_latex_to_mathml(value: str) -> str:
    return latex_to_mathml(value, display=False)


REF_MARK_RE = re.compile(r"(\[(?:\d+(?:-\d+)?)(?:[\uFF0C,]\s*\d+(?:-\d+)?)*\])")


def append_plain_text_with_refs(paragraph, text: str, size=10.5, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", bold=False):
    for chunk in REF_MARK_RE.split(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        set_run_font(run, east_asia, ascii_font, size, bold)
        if REF_MARK_RE.fullmatch(chunk):
            run.font.superscript = True


def append_inline_content(paragraph, text: str, size=10.5, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", bold=False):
    parts = re.split(r"(\$[^$]+\$)", text.replace("`", ""))
    for part in parts:
        if not part:
            continue
        if part.startswith("$") and part.endswith("$"):
            omml = mathml_to_omml(inline_latex_to_mathml(part[1:-1]))
            paragraph._p.append(deepcopy(omml))
        else:
            append_plain_text_with_refs(paragraph, part, size, east_asia, ascii_font, bold)


def add_title(doc, title, author_line="\u4f5c\u8005\u59d3\u540d", affiliation_line="\uff08\u5355\u4f4d\u540d\u79f0    \u90ae\u7f16\uff09"):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(title)
    set_run_font(run, "\u9ed1\u4f53", "SimHei", 18, False)

    blank = doc.add_paragraph()
    blank.alignment = WD_ALIGN_PARAGRAPH.CENTER
    blank.paragraph_format.space_after = Pt(1)
    run = blank.add_run("")
    set_run_font(run, "\u9ed1\u4f53", "SimHei", 9, False)

    for line in (author_line, affiliation_line):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(line)
        set_run_font(run, "\u9ed1\u4f53", "SimHei", 9, False)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5 if level == 1 else 3)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, "\u9ed1\u4f53", "SimHei", 12, False)
    else:
        set_run_font(run, "\u5b8b\u4f53", "Times New Roman", 12, False)
    return paragraph


def add_reference_heading(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("\u53c2\u8003\u6587\u732e")
    set_run_font(run, "\u9ed1\u4f53", "SimHei", 12, False)


def add_body(doc, text, first_indent=True, size=10.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(2)
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(21)

    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            append_inline_content(paragraph, part[2:-2], size=size, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", bold=False)
        else:
            append_inline_content(paragraph, part, size=size, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", bold=False)
    return paragraph


def add_abstract(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(3)
    label = paragraph.add_run("\u6458\u8981\uff1a")
    set_run_font(label, "\u9ed1\u4f53", "SimHei", 12, False)
    append_inline_content(paragraph, text, size=10.5, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", bold=False)


def add_keywords(doc, text):
    text = text.replace("**", "").strip()
    for prefix in ("\u5173\u952e\u8bcd\uff1a", "\u5173\u952e\u8bcd:"):
        if text.startswith(prefix):
            text = text[len(prefix):].strip()
            break
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    label = paragraph.add_run("\u5173\u952e\u8bcd\uff1a")
    set_run_font(label, "\u9ed1\u4f53", "SimHei", 12, False)
    append_inline_content(paragraph, text, size=10.5, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", bold=False)


def add_caption(doc, text, is_table=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text.strip("*"))
    set_run_font(run, "\u9ed1\u4f53" if is_table else "\u5b8b\u4f53", "SimHei" if is_table else "Times New Roman", 9, False)


def add_image(doc, relative_path, alt_text, width_cm=15.8):
    image_path = (ROOT / relative_path).resolve()
    if not image_path.exists():
        add_body(doc, f"[\u7f3a\u5c11\u56fe\u7247\uff1a{relative_path}]", first_indent=False)
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("name", alt_text)
            doc_pr.set("descr", alt_text)


def add_screenshot_placeholder(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.8)
    cell = table.cell(0, 0)
    cell.width = Cm(15.8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F3F5F7")
    set_table_borders(table, color="AAB2BA", size="6")
    row = table.rows[0]
    row.height = Cm(3.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("\u6b64\u5904\u63d2\u5165\u56fe2\u7cfb\u7edf\u754c\u9762\u622a\u56fe\n\u5efa\u8bae\u4e0a\u4e0b\u7ec4\u5408\u5355\u822a\u6b21\u5730\u56fe\u754c\u9762\u4e0e\u5e74\u5ea6\u8bc4\u4f30\u754c\u9762")
    set_run_font(run, "\u5b8b\u4f53", "Times New Roman", 9, False)


def normalize_equation(eq: str) -> str:
    eq = " ".join(line.strip() for line in eq.splitlines() if line.strip())
    eq = eq.replace(r"\begin{aligned}", r"\begin{aligned} ").replace(r"\end{aligned}", r" \end{aligned}")
    eq = re.sub(r"\s+", " ", eq).strip()
    # Match current Markdown quirks around line breaks and braces.
    eq = eq.replace(r"\mathrm{SFOC}_{\mathrm{main} }", r"\mathrm{SFOC}_{\mathrm{main} }")
    return eq


def strip_mathml_equation_number(mathml: str) -> str:
    return re.sub(r'\s*<mspace\s+width="2em"\s*/>\s*<mtext>\(\d+\)</mtext>', '', mathml)


def add_seq_field(paragraph, name="Equation", result_text="1"):
    begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin_run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {name} \\* ARABIC "
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    result_run = paragraph.add_run(result_text)
    set_run_font(result_run, "\u5b8b\u4f53", "Times New Roman", 10.5, False)

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)



def add_math(doc, equation: str):
    key = normalize_equation(equation)
    body, tag = split_equation_tag(key)
    mathml = latex_to_mathml(body, display=True)

    transform = etree.XSLT(etree.parse(str(MML2OMML)))
    omml = transform(etree.fromstring(mathml.encode("utf-8"))).getroot()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2)
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(9.0), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Cm(18.0), WD_TAB_ALIGNMENT.RIGHT)

    paragraph.add_run("\t")
    paragraph._p.append(deepcopy(omml))
    paragraph.add_run("\t(")
    add_seq_field(paragraph, "Equation", tag or "1")
    paragraph.add_run(")")


def visual_len(text: str) -> int:
    return sum(2 if ord(ch) > 127 else 1 for ch in text)


def content_based_widths(rows, total_cm=15.8, min_cm=1.6):
    col_count = max(len(row) for row in rows)
    weights = []
    for col_index in range(col_count):
        max_len = max((visual_len(row[col_index]) for row in rows if col_index < len(row)), default=1)
        weights.append(max(max_len, 6))
    total_weight = sum(weights) or 1
    widths = [max(min_cm, total_cm * weight / total_weight) for weight in weights]
    scale = total_cm / sum(widths)
    return [width * scale for width in widths]


def add_markdown_table(doc, raw_rows):
    parsed = []
    for raw in raw_rows:
        cells = [clean_inline_math(cell.strip()) for cell in raw.strip().strip("|").split("|")]
        if cells and all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return

    widths = content_based_widths(parsed)
    table = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col_index, width in enumerate(widths):
        table.columns[col_index].width = Cm(width)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for row_index, row in enumerate(parsed):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if col_index < len(widths):
                cell.width = Cm(widths[col_index])
            if row_index == 0:
                set_cell_shading(cell, "E8EDF2")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, "\u5b8b\u4f53", "Times New Roman", 7.5, False)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)



def add_reference(doc, text):
    text = re.sub(r"^\[\d+\]\s*", "", text.strip())
    paragraph = doc.add_paragraph()
    apply_reference_numbering(paragraph, ensure_reference_numbering(doc))
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(1)
    append_inline_content(paragraph, text, size=10.5, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", bold=False)


def parse_markdown(doc, markdown, author_line="\u4f5c\u8005\u59d3\u540d", affiliation_line="\uff08\u5355\u4f4d\u540d\u79f0    \u90ae\u7f16\uff09"):
    lines = markdown.splitlines()
    add_title(doc, lines[0].lstrip("# ").strip(), author_line, affiliation_line)

    index = 1
    table_lines = []
    pending_table_caption = None
    pending_figure_caption = None
    abstract_pending = False
    math_lines = []
    in_math = False

    def flush_table():
        nonlocal table_lines, pending_table_caption
        if table_lines:
            if pending_table_caption:
                add_caption(doc, pending_table_caption, is_table=True)
                pending_table_caption = None
            add_markdown_table(doc, table_lines)
            table_lines = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped == "$$":
            flush_table()
            if in_math:
                add_math(doc, "\n".join(math_lines))
                math_lines = []
                in_math = False
            else:
                in_math = True
            index += 1
            continue

        if in_math:
            math_lines.append(raw)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines.append(stripped)
            index += 1
            continue
        flush_table()

        if not stripped:
            index += 1
            continue

        clean = stripped.replace("**", "")
        if stripped == "## \u6458\u8981":
            abstract_pending = True
        elif abstract_pending:
            add_abstract(doc, stripped)
            abstract_pending = False
        elif clean.startswith("\u5173\u952e\u8bcd\uff1a") or clean.startswith("\u5173\u952e\u8bcd:"):
            add_keywords(doc, stripped)
        elif stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            if heading_text == "\u53c2\u8003\u6587\u732e":
                add_reference_heading(doc)
            else:
                add_heading(doc, heading_text, 1)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
        elif stripped.startswith("**\u56fe"):
            pending_figure_caption = stripped
        elif stripped.startswith("**\u8868"):
            pending_table_caption = stripped
        elif stripped.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if match:
                add_image(doc, match.group(2), match.group(1))
                if pending_figure_caption:
                    add_caption(doc, pending_figure_caption, is_table=False)
                    pending_figure_caption = None
        elif stripped.startswith("> \u6b64\u5904\u63d2\u5165"):
            add_screenshot_placeholder(doc)
            if pending_figure_caption:
                add_caption(doc, pending_figure_caption, is_table=False)
                pending_figure_caption = None
            while index + 1 < len(lines) and lines[index + 1].strip().startswith(">"):
                index += 1
        elif stripped.startswith("[") and re.match(r"^\[\d+\]", stripped):
            add_reference(doc, stripped)
        else:
            add_body(doc, stripped, first_indent=True)
        index += 1

    flush_table()
    if pending_figure_caption:
        add_caption(doc, pending_figure_caption, is_table=False)


def main():
    global ROOT, MML2OMML, PANDOC, EQUATION_MAP

    parser = argparse.ArgumentParser(description="Convert a Markdown paper to DOCX with editable Word equations.")
    parser.add_argument("--input", required=True, type=Path, help="Source Markdown file encoded as UTF-8.")
    parser.add_argument("--output", type=Path, help="Output DOCX path. Defaults to input filename with .docx suffix.")
    parser.add_argument("--root", type=Path, help="Asset root for Markdown images. Defaults to the input file directory.")
    parser.add_argument("--mml2omml", type=Path, default=MML2OMML, help="Path to Microsoft Office MML2OMML.XSL.")
    parser.add_argument("--pandoc", default=PANDOC, help="Pandoc executable used for generic LaTeX-to-MathML conversion.")
    parser.add_argument("--equation-map", type=Path, help="Optional JSON object mapping LaTeX expressions to complete MathML strings.")
    parser.add_argument("--author-line", default="\u4f5c\u8005\u59d3\u540d", help="Author line inserted below the title.")
    parser.add_argument("--affiliation-line", default="\uff08\u5355\u4f4d\u540d\u79f0    \u90ae\u7f16\uff09", help="Affiliation line inserted below the author.")
    args = parser.parse_args()

    source = args.input.resolve()
    output = (args.output or source.with_suffix(".docx")).resolve()
    ROOT = (args.root.resolve() if args.root else source.parent)
    MML2OMML = args.mml2omml.resolve()
    PANDOC = args.pandoc
    EQUATION_MAP = load_equation_map(args.equation_map.resolve() if args.equation_map else None)

    if not source.exists():
        raise FileNotFoundError(source)
    if not MML2OMML.exists():
        raise FileNotFoundError(f"Cannot find Office MathML-to-OMML transform: {MML2OMML}")

    markdown = source.read_text(encoding="utf-8")
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.2)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    parse_markdown(document, markdown, args.author_line, args.affiliation_line)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    repair_docx_xml(output)
    stats = inspect_docx(output)
    print(f"output={output}")
    print(f"omath_count={stats['omath_count']}")
    print(f"empty_nary_count={stats['empty_nary_count']}")
    print(f"raw_dollar_count={stats['raw_dollar_count']}")


def move_child_before(parent, child, before_tags):
    if child is None:
        return
    parent.remove(child)
    for existing in list(parent):
        if existing.tag in before_tags:
            existing.addprevious(child)
            return
    parent.append(child)


def repair_docx_xml(path: Path):
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    ns = {"w": w_ns, "m": m_ns}
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(path, "r") as zf:
            zf.extractall(tmp_path)

        document_xml = tmp_path / "word" / "document.xml"
        tree = etree.parse(str(document_xml))
        for tbl_pr in tree.xpath("//w:tblPr", namespaces=ns):
            tbl_borders = tbl_pr.find(f"{{{w_ns}}}tblBorders")
            move_child_before(tbl_pr, tbl_borders, {f"{{{w_ns}}}shd", f"{{{w_ns}}}tblLayout", f"{{{w_ns}}}tblCellMar", f"{{{w_ns}}}tblLook"})
        for tc_pr in tree.xpath("//w:tcPr", namespaces=ns):
            shd = tc_pr.find(f"{{{w_ns}}}shd")
            move_child_before(tc_pr, shd, {f"{{{w_ns}}}noWrap", f"{{{w_ns}}}tcMar", f"{{{w_ns}}}textDirection", f"{{{w_ns}}}tcFitText", f"{{{w_ns}}}vAlign", f"{{{w_ns}}}hideMark"})
        for nary in tree.xpath("//m:nary[m:e[not(.//m:t)]]", namespaces=ns):
            parent = nary.getparent()
            if parent is None:
                continue
            idx = parent.index(nary)
            if idx + 1 >= len(parent):
                continue
            next_el = parent[idx + 1]
            e_el = nary.find(f"{{{m_ns}}}e")
            if e_el is None:
                continue
            parent.remove(next_el)
            e_el.append(next_el)
        tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone="yes")

        settings_xml = tmp_path / "word" / "settings.xml"
        if settings_xml.exists():
            settings_tree = etree.parse(str(settings_xml))
            for zoom in settings_tree.xpath("//w:zoom", namespaces=ns):
                if zoom.get(f"{{{w_ns}}}percent") is None:
                    zoom.set(f"{{{w_ns}}}percent", "100")
            settings_tree.write(str(settings_xml), encoding="UTF-8", xml_declaration=True, standalone="yes")

        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for item in tmp_path.rglob("*"):
                if item.is_file():
                    zf.write(item, item.relative_to(tmp_path).as_posix())


def inspect_docx(path: Path) -> dict:
    m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    ns = {"m": m_ns}
    with zipfile.ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    return {
        "omath_count": len(tree.xpath("//m:oMath | //m:oMathPara", namespaces=ns)),
        "empty_nary_count": len(tree.xpath("//m:nary[m:e[not(.//m:t)]]", namespaces=ns)),
        "raw_dollar_count": xml.count(b"$"),
    }

if __name__ == "__main__":
    main()
